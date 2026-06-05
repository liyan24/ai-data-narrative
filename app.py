import streamlit as st
from streamlit.components.v1 import html as st_html
import sys
from pathlib import Path
import pandas as pd
import json
import io
from datetime import datetime

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from src.pipeline import DataNarrativePipeline
from src.config import OUTPUT_DIR
from src.data_schema import TypeRegistry

st.set_page_config(
    page_title="AI 数据叙事系统",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Custom CSS
st.markdown("""
<style>
    .main-title { font-size: 2.2rem; font-weight: 700; color: #1f77b4; margin-bottom: 0.5rem; }
    .subtitle { font-size: 1.1rem; color: #666; margin-bottom: 2rem; }
    .chart-card { 
        background: #f8f9fa; border-radius: 12px; padding: 20px; 
        margin-bottom: 24px; border: 1px solid #e9ecef;
    }
    .chart-title { font-size: 1.3rem; font-weight: 600; color: #333; margin-bottom: 8px; }
    .desc-text { font-size: 0.95rem; color: #495057; line-height: 1.6; margin-bottom: 6px; }
    .meaning-text { font-size: 0.95rem; color: #0c5460; line-height: 1.6; 
                    background: #d1ecf1; padding: 12px; border-radius: 8px; 
                    border-left: 4px solid #17a2b8; }
    .story-section { 
        background: #fff3cd; border-radius: 12px; padding: 20px; 
        margin-bottom: 16px; border-left: 4px solid #ffc107;
    }
    .story-title { font-size: 1.2rem; font-weight: 600; color: #856404; margin-bottom: 8px; }
    .story-content { font-size: 1rem; color: #533f03; line-height: 1.8; }
    .metric-card { 
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white; border-radius: 12px; padding: 20px; text-align: center;
    }
    .metric-value { font-size: 2rem; font-weight: 700; }
    .metric-label { font-size: 0.9rem; opacity: 0.9; }
    .stButton>button { width: 100%; }
</style>
""", unsafe_allow_html=True)


# ─────────── Word/PDF Generator ───────────

def generate_word_report(result: dict) -> bytes:
    """Generate Word document from analysis results"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Title
    title = doc.add_heading('AI 数据叙事报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    phases = result.get("phases", {})
    data_load = phases.get("data_load", {})
    shape = data_load.get("shape", {}) if isinstance(data_load, dict) else {}
    
    p = doc.add_paragraph()
    p.add_run(f'数据源: {data_load.get("source_name", "N/A")}\n').bold = True
    p.add_run(f'维度: {shape.get("rows", "N/A")} 行 x {shape.get("columns", "N/A")} 列\n')
    p.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # User profile
    user_intent = phases.get("user_intent", {})
    if user_intent and isinstance(user_intent, dict):
        up = user_intent.get("user_profile", {})
        if isinstance(up, dict) and up.get("role"):
            doc.add_heading("用户画像", 1)
            doc.add_paragraph(f'角色: {up.get("role", "N/A")}')
            doc.add_paragraph(f'行业: {up.get("industry", "N/A")}')
            doc.add_paragraph(f'目标: {up.get("goal", "N/A")}')
            doc.add_paragraph(f'专业水平: {up.get("expertise_level", "N/A")}')
    
    # Charts
    skill_exec = phases.get("skill_execution", {})
    chart_data = skill_exec.get("chart-generation", {}).get("data", {}) if isinstance(skill_exec, dict) and isinstance(skill_exec.get("chart-generation"), dict) else {}
    charts = chart_data.get("charts", []) if isinstance(chart_data, dict) else []
    
    if charts:
        doc.add_heading("数据可视化", 1)
        for chart in charts:
            doc.add_heading(chart.get("title", "图表"), 2)
            if chart.get("description"):
                doc.add_paragraph(f'描述: {chart["description"]}')
            if chart.get("meaning"):
                doc.add_paragraph(f'意义: {chart["meaning"]}')
            # Add timeseries config if present
            ts = chart.get("timeseries_config")
            if ts and isinstance(ts, dict) and ts.get("agg_level"):
                doc.add_paragraph(f'时序聚合: {ts["agg_level"]}').italic = True
            doc.add_paragraph()  # spacing
    
    # Story
    story_data = skill_exec.get("story-composition", {}).get("data", {}) if isinstance(skill_exec, dict) and isinstance(skill_exec.get("story-composition"), dict) else {}
    sections = story_data.get("sections", []) if isinstance(story_data, dict) else []
    
    if sections:
        doc.add_heading("数据故事", 1)
        for section in sections:
            title = section.get("title", "")
            content = section.get("content", "")
            if title:
                doc.add_heading(title, 2)
            if content:
                doc.add_paragraph(content)
    
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_pdf_report(result: dict) -> bytes:
    """Generate PDF from analysis results"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # Try to register Chinese font
    try:
        pdfmetrics.registerFont(TTFont('SimHei', 'SimHei.ttf'))
        chinese_font = 'SimHei'
    except:
        try:
            pdfmetrics.registerFont(TTFont('MicrosoftYaHei', 'msyh.ttf'))
            chinese_font = 'MicrosoftYaHei'
        except:
            chinese_font = 'Helvetica'  # fallback
    
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.8*inch, bottomMargin=0.8*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'], fontName=chinese_font,
        fontSize=24, alignment=TA_CENTER, spaceAfter=20, textColor='#1f77b4'
    )
    heading_style = ParagraphStyle(
        'CustomHeading', parent=styles['Heading2'], fontName=chinese_font,
        fontSize=16, spaceAfter=12, textColor='#333'
    )
    subheading_style = ParagraphStyle(
        'CustomSub', parent=styles['Heading3'], fontName=chinese_font,
        fontSize=13, spaceAfter=8, textColor='#555'
    )
    body_style = ParagraphStyle(
        'CustomBody', parent=styles['Normal'], fontName=chinese_font,
        fontSize=11, leading=18, spaceAfter=10
    )
    
    story = []
    
    # Title
    story.append(Paragraph("AI 数据叙事报告", title_style))
    story.append(Spacer(1, 20))
    
    # Meta
    phases = result.get("phases", {})
    data_load = phases.get("data_load", {})
    shape = data_load.get("shape", {}) if isinstance(data_load, dict) else {}
    
    meta_text = f"""<b>数据源:</b> {data_load.get("source_name", "N/A")}<br/>
    <b>维度:</b> {shape.get("rows", "N/A")} 行 x {shape.get("columns", "N/A")} 列<br/>
    <b>生成时间:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}"""
    story.append(Paragraph(meta_text, body_style))
    story.append(Spacer(1, 20))
    
    # User profile
    user_intent = phases.get("user_intent", {})
    if user_intent and isinstance(user_intent, dict):
        up = user_intent.get("user_profile", {})
        if isinstance(up, dict) and up.get("role"):
            story.append(Paragraph("用户画像", heading_style))
            story.append(Paragraph(f'角色: {up.get("role", "N/A")}', body_style))
            story.append(Paragraph(f'行业: {up.get("industry", "N/A")}', body_style))
            story.append(Paragraph(f'目标: {up.get("goal", "N/A")}', body_style))
            story.append(Spacer(1, 10))
    
    # Charts
    skill_exec = phases.get("skill_execution", {})
    chart_data = skill_exec.get("chart-generation", {}).get("data", {}) if isinstance(skill_exec, dict) and isinstance(skill_exec.get("chart-generation"), dict) else {}
    charts = chart_data.get("charts", []) if isinstance(chart_data, dict) else []
    
    if charts:
        story.append(Paragraph("数据可视化", heading_style))
        for chart in charts:
            story.append(Paragraph(chart.get("title", "图表"), subheading_style))
            if chart.get("description"):
                story.append(Paragraph(f'<b>描述:</b> {chart["description"]}', body_style))
            if chart.get("meaning"):
                story.append(Paragraph(f'<b>意义:</b> {chart["meaning"]}', body_style))
            story.append(Spacer(1, 5))
        story.append(Spacer(1, 10))
    
    # Story sections
    story_data = skill_exec.get("story-composition", {}).get("data", {}) if isinstance(skill_exec, dict) and isinstance(skill_exec.get("story-composition"), dict) else {}
    sections = story_data.get("sections", []) if isinstance(story_data, dict) else []
    
    if sections:
        story.append(Paragraph("数据故事", heading_style))
        for section in sections:
            title = section.get("title", "")
            content = section.get("content", "")
            if title:
                story.append(Paragraph(title, subheading_style))
            if content:
                # Clean markdown for PDF
                content_clean = content.replace('**', '')
                story.append(Paragraph(content_clean, body_style))
            story.append(Spacer(1, 5))
    
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ─────────── App Main ───────────

def main():
    # Sidebar
    with st.sidebar:
        st.markdown("### 📊 AI 数据叙事系统")
        st.markdown("v6.0 — 大模型驱动")
        st.divider()
        
        st.markdown("#### ⚙️ 配置")
        max_charts = st.slider("最大图表数", 1, 10, 5)
        auto_clean = st.checkbox("自动清洗数据", value=True)
        
        st.divider()
        st.markdown("#### 📖 使用说明")
        st.markdown("""
        1. 上传数据文件（CSV/Excel/JSON）
        2. 描述你的分析需求
        3. 点击"生成分析报告"
        4. 查看交互式图表和数据故事
        """)
        
        st.divider()
        st.markdown("#### 📁 支持格式")
        st.markdown("""
        - CSV (.csv)
        - Excel (.xlsx, .xls)
        - JSON (.json, .jsonl)
        - Parquet (.parquet)
        - SQLite (.db, .sqlite)
        """)
    
    # Main content
    st.markdown('<div class="main-title">📊 AI 数据叙事系统</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">上传数据，输入需求，自动生成数据故事</div>', unsafe_allow_html=True)
    
    # Upload section
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("### 📁 数据上传")
        uploaded_file = st.file_uploader(
            "选择数据文件",
            type=["csv", "xlsx", "xls", "json", "jsonl", "parquet", "db", "sqlite"],
            help="支持 CSV、Excel、JSON、JSON Lines、Parquet、SQLite"
        )
    
    with col2:
        st.markdown("### 📝 分析需求")
        user_input = st.text_area(
            "描述你的分析需求（可选）",
            placeholder="例如：我是电商运营，想分析销售趋势和地区分布",
            height=80,
        )
    
    if uploaded_file is not None:
        # Save uploaded file
        data_dir = Path("data")
        data_dir.mkdir(exist_ok=True)
        file_path = data_dir / uploaded_file.name
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        
        # Preview data
        st.divider()
        st.markdown("### 👀 数据预览")
        try:
            if file_path.suffix == ".csv":
                df_preview = pd.read_csv(file_path, nrows=10)
            elif file_path.suffix in [".xlsx", ".xls"]:
                df_preview = pd.read_excel(file_path, nrows=10)
            elif file_path.suffix == ".json":
                df_preview = pd.read_json(file_path, nrows=10)
            elif file_path.suffix == ".jsonl":
                df_preview = pd.read_json(file_path, lines=True, nrows=10)
            elif file_path.suffix == ".parquet":
                df_preview = pd.read_parquet(file_path).head(10)
            else:
                df_preview = None
            
            if df_preview is not None:
                st.dataframe(df_preview, use_container_width=True)
                st.caption(f"共 {len(df_preview.columns)} 列，显示前 {min(10, len(df_preview))} 行")
        except Exception as e:
            st.warning(f"预览加载失败: {e}")
        
        # Run analysis button
        st.divider()
        if st.button("🚀 生成分析报告", type="primary", use_container_width=True):
            # Progress tracking
            progress_container = st.empty()
            status_container = st.empty()
            
            def progress_callback(phase, status, msg):
                # Update in session state for display
                if "progress" not in st.session_state:
                    st.session_state["progress"] = {}
                st.session_state["progress"][phase] = {"status": status, "msg": msg}
            
            # Show progress
            with st.status("🤖 AI 正在分析数据...", expanded=True) as status:
                try:
                    # Phase 0-1 progress
                    st.write("📋 Phase 0: 理解用户需求...")
                    
                    pipeline = DataNarrativePipeline(
                        output_dir=OUTPUT_DIR,
                        max_charts=max_charts,
                        enable_user_intent=True,
                        enable_skill_director=True,
                        verbose=False,
                        auto_clean=auto_clean,
                    )
                    
                    st.write("📁 Phase 1: 加载数据并理解结构...")
                    
                    # Phase 2-3
                    st.write("🎯 Phase 2: 制定执行计划...")
                    st.write("⚙️ Phase 3: 执行技能（质量检查、统计、洞察、图表、叙事）...")
                    
                    result = pipeline.run(
                        str(file_path),
                        user_input=user_input,
                        progress_callback=progress_callback,
                    )
                    
                    st.session_state["result"] = result
                    status.update(label="✅ 分析完成！", state="complete", expanded=False)
                    
                except Exception as e:
                    status.update(label=f"❌ 分析失败: {e}", state="error")
                    import traceback
                    st.error(traceback.format_exc())
    
    # Display results
    if "result" in st.session_state:
        result = st.session_state["result"]
        phases = result.get("phases", {})
        
        st.divider()
        st.markdown("## 📋 分析结果")
        
        # Metrics row
        metric_cols = st.columns(4)
        
        # Data dimensions — FIX: use shape.rows / shape.columns
        data_load = phases.get("data_load", {})
        if isinstance(data_load, dict):
            shape = data_load.get("shape", {}) if isinstance(data_load, dict) else {}
            rows = shape.get("rows", "N/A") if isinstance(shape, dict) else "N/A"
            cols = shape.get("columns", "N/A") if isinstance(shape, dict) else "N/A"
            with metric_cols[0]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{rows}</div>
                    <div class="metric-label">数据行数</div>
                </div>
                """, unsafe_allow_html=True)
            with metric_cols[1]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{cols}</div>
                    <div class="metric-label">数据列数</div>
                </div>
                """, unsafe_allow_html=True)
        
        # Skills count
        skill_execution = phases.get("skill_execution", {})
        if isinstance(skill_execution, dict):
            success_count = sum(1 for sr in skill_execution.values() if sr.get("status") == "success")
            with metric_cols[2]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{success_count}</div>
                    <div class="metric-label">成功技能</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Charts count
            chart_result = skill_execution.get("chart-generation", {})
            chart_data = chart_result.get("data", {}) if isinstance(chart_result, dict) else {}
            chart_count = chart_data.get("count", 0) if isinstance(chart_data, dict) else 0
            with metric_cols[3]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{chart_count}</div>
                    <div class="metric-label">生成图表</div>
                </div>
                """, unsafe_allow_html=True)
        
        # User intent
        user_intent = phases.get("user_intent", {})
        if isinstance(user_intent, dict) and user_intent.get("user_profile"):
            st.divider()
            st.markdown("### 👤 用户画像")
            up = user_intent["user_profile"]
            ui = user_intent.get("user_intent", {})
            cols = st.columns(3)
            with cols[0]:
                st.markdown(f"**角色:** {up.get('role', 'N/A')}")
                st.markdown(f"**行业:** {up.get('industry', 'N/A')}")
            with cols[1]:
                st.markdown(f"**目标:** {up.get('goal', 'N/A')}")
                st.markdown(f"**专业水平:** {up.get('expertise_level', 'N/A')}")
            with cols[2]:
                st.markdown(f"**意图类型:** {ui.get('intent_type', 'N/A')}")
                st.markdown(f"**置信度:** {ui.get('intent_confidence', 'N/A')}")
        
        # Data Understanding — NEW
        data_understanding = phases.get("data_understanding", {})
        if isinstance(data_understanding, dict) and data_understanding.get("business_domain"):
            st.divider()
            st.markdown("### 🧠 数据理解（大模型分析）")
            
            with st.expander("📋 业务理解详情", expanded=True):
                cols = st.columns(3)
                with cols[0]:
                    st.markdown(f"**📁 业务领域:** {data_understanding.get('business_domain', 'N/A')}")
                    st.markdown(f"**📊 业务场景:** {data_understanding.get('business_scenario', 'N/A')}")
                with cols[1]:
                    st.markdown(f"**📈 核心指标:** {', '.join(data_understanding.get('key_metrics', [])[:5])}")
                    st.markdown(f"**📂 核心维度:** {', '.join(data_understanding.get('key_dimensions', [])[:5])}")
                with cols[2]:
                    st.markdown(f"**⏰ 时间列:** {data_understanding.get('time_column', 'N/A')}")
                    st.markdown(f"**🆔 ID列:** {data_understanding.get('id_column', 'N/A')}")
                
                st.markdown(f"**📝 数据描述:** {data_understanding.get('table_description', 'N/A')}")
                
                # 列详情表格
                columns = data_understanding.get("columns", [])
                if columns:
                    df_cols = pd.DataFrame([
                        {
                            "列名": c.get("name", ""),
                            "业务含义": c.get("business_meaning", ""),
                            "类型": c.get("data_type", ""),
                            "角色": c.get("business_role", ""),
                            "建议聚合": ", ".join(c.get("suggested_aggregations", [])),
                        }
                        for c in columns
                    ])
                    st.dataframe(df_cols, use_container_width=True, hide_index=True)
        
        # Charts section — DYNAMIC: only show if there are charts with html
        st.divider()
        st.markdown("## 📈 数据可视化")
        
        chart_result = skill_execution.get("chart-generation", {})
        chart_data = chart_result.get("data", {}) if isinstance(chart_result, dict) else {}
        charts = chart_data.get("charts", []) if isinstance(chart_data, dict) else []
        
        # Filter: only show charts with valid html
        valid_charts = [c for c in charts if c.get("html") and len(c.get("html", "")) > 100]
        
        if valid_charts:
            for i, chart in enumerate(valid_charts):
                with st.container():
                    st.markdown('<div class="chart-card">', unsafe_allow_html=True)
                    
                    # Chart title + timeseries badge
                    title = chart.get("title", f"图表 {i+1}")
                    ts_config = chart.get("timeseries_config")
                    if ts_config and isinstance(ts_config, dict) and ts_config.get("agg_level") and ts_config.get("agg_level") != "raw":
                        title += f" [📅 {ts_config['agg_level']}聚合]"
                    st.markdown(f'<div class="chart-title">{title}</div>', unsafe_allow_html=True)
                    
                    # ECharts HTML
                    html_content = chart.get("html", "")
                    st_html(html_content, height=450, scrolling=True)
                    
                    # Two-sentence description
                    desc = chart.get("description", "")
                    meaning = chart.get("meaning", "")
                    
                    if desc:
                        st.markdown(f'<div class="desc-text">📋 <b>描述：</b>{desc}</div>', unsafe_allow_html=True)
                    if meaning:
                        st.markdown(f'<div class="meaning-text">💡 <b>意义：</b>{meaning}</div>', unsafe_allow_html=True)
                    
                    # Timeseries explanation
                    if ts_config and isinstance(ts_config, dict) and ts_config.get("explanation"):
                        st.info(f"📅 时序分析: {ts_config['explanation']}")
                    
                    st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("暂无可用图表数据")
        
        # Data Story section — DYNAMIC: only show if there are sections
        st.divider()
        st.markdown("## 📖 数据故事")
        
        story_result = skill_execution.get("story-composition", {})
        story_data = story_result.get("data", {}) if isinstance(story_result, dict) else {}
        sections = story_data.get("sections", []) if isinstance(story_data, dict) else []
        
        if sections:
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                
                if title or content:
                    with st.container():
                        st.markdown('<div class="story-section">', unsafe_allow_html=True)
                        if title:
                            st.markdown(f'<div class="story-title">{title}</div>', unsafe_allow_html=True)
                        if content:
                            st.markdown(f'<div class="story-content">{content}</div>', unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("暂无数据故事")
        
        # Download section — Word + PDF
        st.divider()
        st.markdown("### 📥 下载报告")
        
        download_cols = st.columns(3)
        
        with download_cols[0]:
            # HTML report
            report_path = result.get("report_path")
            if report_path and Path(report_path).exists():
                with open(report_path, "rb") as f:
                    st.download_button(
                        label="📄 HTML 报告",
                        data=f.read(),
                        file_name=Path(report_path).name,
                        mime="text/html",
                        use_container_width=True,
                    )
            else:
                st.button("📄 HTML 报告", disabled=True, use_container_width=True)
        
        with download_cols[1]:
            # Word report
            try:
                word_bytes = generate_word_report(result)
                st.download_button(
                    label="📝 Word 报告",
                    data=word_bytes,
                    file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Word生成失败: {e}")
        
        with download_cols[2]:
            # PDF report
            try:
                pdf_bytes = generate_pdf_report(result)
                st.download_button(
                    label="📕 PDF 报告",
                    data=pdf_bytes,
                    file_name=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF生成失败: {e}")
        
        # Performance
        perf = result.get("performance")
        if perf and isinstance(perf, dict):
            st.caption(f"⏱️ 总耗时: {perf.get('total_elapsed_seconds', 'N/A')} 秒")


if __name__ == "__main__":
    main()
